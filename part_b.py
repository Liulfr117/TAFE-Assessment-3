'''
Program Description: This program can create and make user-specified changes to Word files. The available actions are creating, adding headings.
The user is asked which action they would like to perform, and then asked to specify arguments. The program performs the chosen action on the
specified pdf file using functionality from the docx library, and then saves the output file to the specified location.
'''

# Name: Dylan Bailey
# Student ID: 20114978

# Date                   Version                    Reason for Change
# 26/10/2024             1                          Creation of program
# 28/10/2024             2                          Changed heading creation code to allow users to specify what level heading they want, rather than
#                                                   rather than prompt them for two headings sequentially
# 28/10/2024             3                          Changed picture insertion code so that it has exception handling to deal with errors

from docx import Document
from docx.shared import Inches

def main():
    # Initialize by creating a Word document
    document = Document()

    # A while loop is used to create a menu in which the user can specify what action they want to perform. If they choose an action number that
    # doesn't exist, the program will reprompt them.
    while True:
        print("\n1. Add Heading to document")
        print("2. Add Paragraph to a specific heading")
        print("3. Add Picture to document")
        print("4. Retrieve full text from document")
        print("5. Save and Exit")

        choice = input("Choose an option (1-5): ")

        # Conditional that handles adding heading to a document. Asks user for input regarding heading text and level, then adds this information to the created document
        if choice == '1':
            heading_text = input("Enter the heading text: ")
            level = int(input("Enter heading level (1-3): "))
            document.add_heading(heading_text, level=level)
            print(f"Heading '{heading_text}' added to the document.")

        # Conditional statement that handles adding user-iputted paragraphs to a word document
        elif choice == '2':
            paragraph_text = input("Enter the paragraph text: ")
            document.add_paragraph(paragraph_text)
            print(f"Paragraph added: '{paragraph_text}'")

        # Conditional statement that handles inserting an image into a word document
        elif choice == '3':
            image_path = input("Enter the path to the image file: ")
            try:
                document.add_picture(image_path, width=Inches(2))
                print("Image added to the document.")
            except Exception as e:
                print(f"Error adding image: {e}")

        # Conditional statement that handles extracting and displaying the text contained within the document
        elif choice == '4':
            print("Full text from document:")
            print(get_full_text(document))

        # Saves the created and modified document using a user-defined file name
        elif choice == '5':
            save_document(document)
            print("Document saved and program exited.")
            break

        else:
            print("Invalid choice, please try again.")

# Function to retrieve full text from the document
def get_full_text(document):
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

# Function to save the document
def save_document(document):
    output_file = input("Enter the output document name (with .docx extension): ")
    document.save(output_file)
    print(f"Document saved as '{output_file}'")


if __name__ == "__main__":
    main()
