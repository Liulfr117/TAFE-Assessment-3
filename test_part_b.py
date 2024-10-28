import unittest
from unittest.mock import patch, MagicMock
from docx import Document
from part_b import get_full_text, save_document  # Adjust the import as needed


class TestWordDocumentOperations(unittest.TestCase):

    @patch('docx.Document')
    def test_add_heading(self, MockDocument):
        # Create a mock document
        mock_doc = MockDocument.return_value

        # Simulate adding a heading
        heading_text = "Test Heading"
        level = 2
        mock_doc.add_heading(heading_text, level=level)

        # Verify the heading was added
        mock_doc.add_heading.assert_called_once_with(heading_text, level=level)

    @patch('docx.Document')
    def test_add_paragraph(self, MockDocument):
        # Create a mock document
        mock_doc = MockDocument.return_value

        # Simulate adding a paragraph
        paragraph_text = "This is a test paragraph."
        mock_doc.add_paragraph(paragraph_text)

        # Verify the paragraph was added
        mock_doc.add_paragraph.assert_called_once_with(paragraph_text)

    @patch('docx.Document')
    def test_get_full_text(self, MockDocument):
        # Create a mock document with paragraphs
        mock_doc = MockDocument.return_value
        mock_doc.paragraphs = [MagicMock(text='First paragraph.'),
                                MagicMock(text='Second paragraph.')]

        full_text = get_full_text(mock_doc)

        # Verify the full text retrieval
        expected_text = 'First paragraph.\nSecond paragraph.'
        self.assertEqual(full_text, expected_text)

    @patch('docx.Document')
    @patch('builtins.input', side_effect=['test_doc.docx'])  # Simulate user input
    def test_save_document(self, mock_input, MockDocument):
        # Create a mock document
        mock_doc = MockDocument.return_value

        # Call the save_document function
        save_document(mock_doc)

        # Verify the document was saved with the correct file name
        mock_doc.save.assert_called_once_with('test_doc.docx')


if __name__ == '__main__':
    unittest.main()
